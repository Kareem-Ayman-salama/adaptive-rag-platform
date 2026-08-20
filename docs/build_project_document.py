"""Build Word documentation for DocuMind AI from the Markdown source."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "PROJECT_DOCUMENTATION.md"
OUT = ROOT / "docs" / "DocuMind_AI_Project_Documentation.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(88, 96, 112)
LIGHT_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    """Apply background shading to a table cell."""

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table) -> None:
    """Apply compact readable cell margins."""

    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side in ("top", "bottom", "start", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), "120" if side in {"start", "end"} else "80")
        node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    """Configure document geometry and styles."""

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.text = "DocuMind AI Project Documentation"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = MUTED


def add_cover(doc: Document) -> None:
    """Add a simple first page cover."""

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DocuMind AI")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        "Adaptive Multimodal RAG for Source-Grounded PDF Intelligence"
    )
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = MUTED

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    set_cell_margins(meta)
    rows = [
        ("Supervisor", "Eng / Jana Hatem"),
        ("Assistant Engineer", "Gad Amr"),
        ("Team", "Kemet AI"),
        ("Deliverables", "Documentation, Presentation, Web Application"),
    ]
    for row, (key, value) in zip(meta.rows, rows, strict=True):
        row.cells[0].text = key
        row.cells[1].text = value
        set_cell_shading(row.cells[0], LIGHT_FILL)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()
    lead = doc.add_paragraph()
    lead.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead_run = lead.add_run(
        "Built around trust, evidence, hallucination control, and education."
    )
    lead_run.font.size = Pt(12)
    lead_run.font.bold = True

    doc.add_section(WD_SECTION.NEW_PAGE)


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    """Convert a simple Markdown table to a Word table."""

    rows = []
    for line in lines:
        cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    set_cell_margins(table)

    for row_index, row_cells in enumerate(rows):
        for col_index, value in enumerate(row_cells):
            cell = table.cell(row_index, col_index)
            cell.text = value
            if row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if row_index == 0:
                        run.font.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    """Add a compact code block using monospaced text."""

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(55, 65, 81)


def add_markdown_content(doc: Document, text: str) -> None:
    """Parse the project Markdown into a polished Word document."""

    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            heading = stripped.removeprefix("# ").strip()
            if heading != "DocuMind AI - Project Documentation":
                doc.add_heading(heading, level=1)
            i += 1
            continue
        if stripped.startswith("## "):
            doc.add_heading(stripped.removeprefix("## ").strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped.removeprefix("### ").strip(), level=2)
            i += 1
            continue

        if stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.add_run(stripped.removeprefix("- ").strip())
            i += 1
            continue

        if stripped[0].isdigit() and ". " in stripped[:4]:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.add_run(stripped.split(". ", 1)[1].strip())
            i += 1
            continue

        cleaned = stripped.replace("**", "").replace("`", "")
        doc.add_paragraph(cleaned)
        i += 1


def main() -> None:
    """Build the Word documentation file."""

    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_markdown_content(doc, SOURCE.read_text(encoding="utf-8"))
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
